"""
make_doc_s23.py - Génère le document Word pédagogique sur les Séances 2 et 3 du TP.
Même style et même mise en forme que Seance1_Pedagogique.docx.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Couleurs ────────────────────────────────────────────────────────────
ROUGE_CNAM = RGBColor(0xC5, 0x0B, 0x26)
BLEU_TITRE = RGBColor(0x1F, 0x49, 0x7D)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)


# ── Helpers (identiques au document de la Séance 1) ──────────────────────
def set_run_font(run, size=11, bold=False, italic=False, color=None, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def ajouter_titre_principal(texte):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(texte), size=22, bold=True, color=ROUGE_CNAM)


def ajouter_sous_titre(texte):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(texte), size=13, italic=True, color=BLEU_TITRE)


def h1(texte):
    p = doc.add_paragraph()
    set_run_font(p.add_run(texte), size=15, bold=True, color=ROUGE_CNAM)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), "C50B26")
    pBdr.append(bottom); pPr.append(pBdr)


def h2(texte):
    p = doc.add_paragraph()
    set_run_font(p.add_run(texte), size=12, bold=True, color=BLEU_TITRE)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def corps(texte, italic=False):
    p = doc.add_paragraph()
    set_run_font(p.add_run(texte), size=11, italic=italic)
    p.paragraph_format.space_after = Pt(4)


def puce(texte, bold_debut=None, indent=0.8):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    if bold_debut:
        set_run_font(p.add_run(bold_debut), size=11, bold=True)
    set_run_font(p.add_run(texte), size=11)


def _shade(cell, hexa):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexa)
    tcPr.append(shd)


def encadre(texte, titre=None, couleur_fond=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    if couleur_fond:
        r, g, b = couleur_fond
        _shade(cell, f"{r:02X}{g:02X}{b:02X}")
    if titre:
        tp = cell.add_paragraph()
        set_run_font(tp.add_run(titre), size=11, bold=True, color=ROUGE_CNAM)
        tp.paragraph_format.space_after = Pt(3)
    cp = cell.add_paragraph()
    set_run_font(cp.add_run(texte), size=11)
    cp.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def code_bloc(lignes):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _shade(cell, "F2F2F2")
    for ligne in lignes:
        cp = cell.add_paragraph()
        set_run_font(cp.add_run(ligne), size=10, name="Courier New")
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()


def tableau(entetes, lignes, mono_cols=(), surligne_idx=None, tailles=None):
    """Table Grid avec en-tête rouge. mono_cols = indices en Courier.
    surligne_idx = index de ligne (dans lignes) à surligner en vert."""
    n = len(entetes)
    t = doc.add_table(rows=1 + len(lignes), cols=n)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # en-tête
    for j, e in enumerate(entetes):
        cell = t.rows[0].cells[j]
        r = cell.paragraphs[0].add_run(e)
        set_run_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade(cell, "C50B26")
    # lignes
    for i, ligne in enumerate(lignes):
        for j, val in enumerate(ligne):
            cell = t.rows[i + 1].cells[j]
            name = "Courier New" if j in mono_cols else "Calibri"
            set_run_font(cell.paragraphs[0].add_run(str(val)), size=10.5, name=name)
        if surligne_idx is not None and i == surligne_idx:
            for cell in t.rows[i + 1].cells:
                _shade(cell, "E2EFDA")
    doc.add_paragraph()


def saut():
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ════════════════════════════════════════════════════════════════════════
saut(); saut()
ajouter_titre_principal("Simulateur à événements discrets")
ajouter_sous_titre("Document pédagogique — Séances 2 et 3")
saut()
ajouter_sous_titre("RCP103 / NCA — CNAM")
ajouter_sous_titre("Pedro Braconnot Velloso & Abdelhak Hidouri")
saut()
encadre(
    "Ce document fait suite au document de la Séance 1. Il explique, de façon "
    "progressive et illustrée, les Séances 2 et 3 du TP : les composants actifs "
    "du système (Client, Queue, Server), puis l'assemblage final (Gateway, "
    "Engine) et les vraies simulations. Après lecture, vous comprendrez le rôle "
    "de chaque classe, le fonctionnement de la boucle de simulation, et la façon "
    "de mesurer et de vérifier les résultats.",
    titre="📘  À propos de ce document",
    couleur_fond=(0xEB, 0xF3, 0xFF))

# ── Continuité avec la Séance 1 ──
h1("Rappel : ce que la Séance 1 a posé")
corps(
    "En Séance 1, on a construit les briques de base du simulateur : la classe "
    "Message (ce qu'on transporte), la classe Event (ce qui se produit à un "
    "instant donné), et le Scheduler (l'ordonnanceur qui garde les événements "
    "triés par ordre chronologique, via un tas binaire, sans sort()). On a aussi "
    "mis en place la trace (journal CSV des événements) et la méthode de test : "
    "une fonction de test par classe, en développement incrémental.")
encadre(
    "Ces briques ne « font » encore rien toutes seules. Les Séances 2 et 3 ajoutent "
    "les composants qui produisent et traitent les messages (Séance 2), puis le "
    "moteur qui orchestre tout et lance les simulations (Séance 3).",
    titre="➡️  Là où on en est",
    couleur_fond=(0xFF, 0xF8, 0xE1))

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  SÉANCE 2
# ════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("SÉANCE 2"), size=20, bold=True, color=ROUGE_CNAM)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p2.add_run("Les composants actifs : Client, Queue, Server"),
             size=13, italic=True, color=BLEU_TITRE)
saut()

# 1. Objectif
h1("1.  Objectif de la Séance 2")
corps(
    "Après les briques passives de la Séance 1, on construit les trois composants "
    "« actifs » du système, ceux qui ont un comportement propre :")
puce(" — il fabrique les messages et décide quand ils arrivent.", bold_debut="le Client (la source)")
puce(" — elle range les messages qui doivent patienter.", bold_debut="la Queue (la file d'attente)")
puce(" — il traite les messages, un à la fois.", bold_debut="le Server (le serveur)")
corps(
    "Ces trois classes ne sont pas encore reliées entre elles (ce sera le travail "
    "de la Gateway en Séance 3). On les écrit et on les teste séparément, toujours "
    "en développement incrémental.")

# 2. Client
h1("2.  La classe Client (générateur de messages)")
h2("2.1  Son rôle")
corps(
    "Le Client représente le dispositif mobile IoT qui envoie des messages vers la "
    "passerelle. Son travail : produire un flot de messages dans le temps. La "
    "question centrale est : à quel rythme les messages arrivent-ils ?")

h2("2.2  Les inter-arrivées exponentielles")
corps(
    "On ne fixe pas un message « toutes les X secondes » de façon régulière : dans "
    "la réalité, les arrivées sont irrégulières. On modélise donc le temps qui "
    "sépare deux arrivées (l'« inter-arrivée ») par une variable aléatoire suivant "
    "une loi exponentielle.")
encadre(
    "λ (lambda) est un TAUX d'arrivée : le nombre moyen de messages par seconde. "
    "Si λ = 4 msg/s, alors le temps moyen entre deux messages est 1/λ = 1/4 = "
    "0,25 s. Attention : moyenne = 1/λ, et non λ. C'est un point qui prête à "
    "confusion (le cours note la formule -λ·ln(U), mais c'est bien 1/λ qui est la "
    "moyenne physiquement correcte — vérifié plus loin).",
    titre="💡  λ est un taux, la moyenne est 1/λ",
    couleur_fond=(0xFF, 0xF8, 0xE1))
corps(
    "Concrètement, pour tirer une inter-arrivée, on utilise la « transformée "
    "inverse » : si U est un nombre aléatoire uniforme dans [0,1[, alors "
    "-ln(1-U)/λ suit une exponentielle de moyenne 1/λ.")
code_bloc([
    "public double nextDelay() {",
    "    double u = rng.nextDouble();        // U dans [0, 1[",
    "    return -Math.log(1.0 - u) / lambda; // moyenne = 1/lambda",
    "}",
])

h2("2.3  Création des messages et événements SEND_MSG")
corps(
    "Le Client crée aussi les Messages (avec un identifiant qui s'incrémente) via "
    "generateMessage(). Dans la simulation, chaque envoi correspond à un événement "
    "SEND_MSG : le Client émet un message à l'instant t, et programme son prochain "
    "envoi à t + nextDelay(). C'est ce mécanisme qui entretient le flot d'arrivées.")

h2("2.4  Méthodes principales")
tableau(
    ["Méthode", "Rôle"],
    [["nextDelay()", "tire une inter-arrivée exponentielle de moyenne 1/λ"],
     ["generateMessage()", "crée le prochain message (IDs incrémentaux)"],
     ["printClient()", "affiche la configuration du client"]],
    mono_cols=(0,))

# 3. Queue
h1("3.  La classe Queue (file d'attente)")
h2("3.1  Son rôle et le principe FIFO")
corps(
    "Quand un message arrive mais que le serveur est occupé, il ne peut pas être "
    "traité tout de suite : il attend dans la file. La Queue gère cette salle "
    "d'attente selon le principe FIFO (First In, First Out) : premier arrivé, "
    "premier servi. C'est le fonctionnement le plus naturel et le plus équitable.")
code_bloc([
    "Arrivées :   #1  #2  #3   ──►   [ #1 | #2 | #3 ]   ──►   service de #1 d'abord",
    "                                 ▲ tête        ▲ queue",
])

h2("3.2  Capacité : infinie ou limitée à K")
corps(
    "La file peut être de capacité infinie (modèle M/M/1 : on accepte toujours), "
    "ou limitée à K places (modèles « bornés » M/M/1/K). La capacité K désigne le "
    "nombre maximal de messages dans tout le système : ceux qui attendent en file "
    "PLUS celui (ou ceux) en cours de service.")
encadre(
    "Quand le système est plein (déjà K messages), un nouveau message arrivant est "
    "REJETÉ (perdu). C'est un compromis : on borne l'attente et on garantit la "
    "stabilité, mais au prix de pertes. Sans borne, la file peut grossir sans fin "
    "si les arrivées dépassent la capacité de traitement.",
    titre="⚠️  File pleine = rejet",
    couleur_fond=(0xFD, 0xEC, 0xEA))
tableau(
    ["Méthode", "Rôle"],
    [["enqueue(msg)", "ajoute un message en queue de file"],
     ["dequeue()", "retire et retourne le message en tête (le plus ancien)"],
     ["isEmpty()", "indique si la file est vide"],
     ["size()", "nombre de messages en attente"]],
    mono_cols=(0,))

# 4. Server
h1("4.  La classe Server (serveur)")
h2("4.1  Son rôle et son état")
corps(
    "Le Server traite les messages, un seul à la fois. Il a donc un état simple : "
    "libre ou occupé. Quand il prend un message, il devient occupé ; quand il a "
    "fini, il redevient libre et peut prendre le suivant.")

h2("4.2  Le temps de service exponentiel")
corps(
    "Comme pour les arrivées, le temps que met le serveur à traiter un message est "
    "aléatoire et suit une loi exponentielle.")
encadre(
    "µ (mu) est le TAUX de service : le nombre moyen de messages traités par "
    "seconde quand le serveur travaille. Si µ = 8 msg/s, le temps de service moyen "
    "est 1/µ = 1/8 = 0,125 s. Là encore : moyenne = 1/µ.",
    titre="💡  µ est un taux, le temps moyen de service est 1/µ",
    couleur_fond=(0xFF, 0xF8, 0xE1))

h2("4.3  Déclenchement de la fin de service (MSG_DEPT)")
corps(
    "Quand un serveur commence à traiter un message à l'instant t, on programme "
    "immédiatement l'événement MSG_DEPT (départ) à t + serviceTime(). Cet "
    "événement marquera la fin du traitement : le serveur se libère et, s'il y a "
    "des messages en file, il sert le suivant.")
tableau(
    ["Méthode", "Rôle"],
    [["serviceTime()", "tire un temps de service exponentiel de moyenne 1/µ"],
     ["occupy() / release()", "passe le serveur occupé / libre"],
     ["isFree()", "indique si le serveur est disponible"],
     ["printServer()", "affiche l'état du serveur"]],
    mono_cols=(0,))

# 5. Pourquoi exponentielles
h1("5.  Pourquoi des lois exponentielles ?")
corps(
    "Le choix de la loi exponentielle n'est pas un hasard : c'est la seule loi "
    "« sans mémoire ». Cela signifie que le temps déjà écoulé n'influence pas le "
    "temps restant.")
encadre(
    "Exemple intuitif : si le temps de service est exponentiel et que le serveur "
    "travaille depuis 0,1 s, la probabilité qu'il finisse dans la prochaine "
    "milliseconde est la même que s'il venait de commencer. Le système « oublie » "
    "son passé. Cette propriété rend les calculs possibles et définit les modèles "
    "dits markoviens — le « M » de M/M/1 (arrivées Markoviennes / service "
    "Markovien / 1 serveur).",
    titre="🧠  La propriété « sans mémoire »",
    couleur_fond=(0xEB, 0xF3, 0xFF))
corps(
    "Vérification dans notre code : en tirant 200 000 inter-arrivées et 200 000 "
    "temps de service, on retrouve bien les moyennes théoriques, ce qui confirme "
    "que λ et µ sont traités comme des taux (moyenne 1/λ et 1/µ) :")
tableau(
    ["Grandeur", "Moyenne empirique (sim.)", "Théorie (1/taux)"],
    [["Inter-arrivées (λ=4)", "0,2503 s", "0,2500 s"],
     ["Temps de service (µ=8)", "0,1254 s", "0,1250 s"]])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  SÉANCE 3
# ════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("SÉANCE 3"), size=20, bold=True, color=ROUGE_CNAM)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p2.add_run("L'assemblage : Gateway, Engine et les simulations"),
             size=13, italic=True, color=BLEU_TITRE)
saut()

# 1. Objectif
h1("1.  Objectif de la Séance 3")
corps(
    "C'est la séance qui réunit tout. On relie les briques (Séance 1) et les "
    "composants actifs (Séance 2) dans un simulateur complet, puis on lance de "
    "vraies simulations pour mesurer les performances et les comparer à la théorie. "
    "Deux nouvelles classes : la Gateway (le système) et l'Engine (le moteur).")

# 2. Gateway
h1("2.  La classe Gateway (la passerelle = le système)")
h2("2.1  Son rôle")
corps(
    "La Gateway est « le système » que l'on étudie. Elle regroupe la file (Queue) "
    "et le ou les serveurs (Server), et applique la logique de décision à chaque "
    "événement. Selon le modèle, elle contient 1 serveur (M/M/1, M/M/1/K) ou 3 "
    "serveurs (M/M/3/8), et une capacité infinie ou limitée à K.")

h2("2.2  Logique de l'événement RECV_MSG (un message arrive)")
corps("Quand un message arrive à la passerelle, trois cas possibles :")
puce(" → le service commence tout de suite, on programme son MSG_DEPT.", bold_debut="un serveur est libre")
puce(" → le message est mis en file d'attente.", bold_debut="tous les serveurs sont occupés (et système non plein)")
puce(" → le message est rejeté (compteur de rejets +1).", bold_debut="le système est plein (cas /K)")
code_bloc([
    "RECV_MSG :",
    "   si système plein        →  REJET",
    "   sinon si serveur libre   →  SERVICE  (programmer MSG_DEPT)",
    "   sinon                    →  MISE EN FILE",
])

h2("2.3  Logique de l'événement MSG_DEPT (fin de service)")
corps(
    "Quand un serveur termine, la Gateway le libère et comptabilise la sortie. "
    "Puis, si la file n'est pas vide, elle sort le message suivant et le met "
    "aussitôt en service (nouveau MSG_DEPT). C'est ce qui fait « tourner » la file.")
code_bloc([
    "MSG_DEPT :",
    "   libérer le serveur ; compter le départ",
    "   si la file n'est pas vide  →  servir le message suivant",
])

# 3. Engine
h1("3.  La classe Engine (moteur de simulation)")
h2("3.1  Le rôle du moteur")
corps(
    "L'Engine contient le main et pilote toute la simulation. Il crée les clients, "
    "construit la passerelle, et fait tourner la boucle à événements discrets "
    "jusqu'à la fin de l'horizon de simulation T.")

h2("3.2  La boucle à événements discrets")
corps(
    "C'est le cœur du simulateur. Tant qu'il reste des événements (et que t < T), "
    "on répète quatre étapes :")
code_bloc([
    "  TANT QUE le scheduler a des événements ET t < T :",
    "    1. extraire l'événement le plus proche   e = scheduler.getEvent()",
    "    2. avancer le temps                       t = e.getEventTime()",
    "    3. mettre à jour l'état + les métriques   (accumuler N·Δt)",
    "    4. traiter l'événement et programmer       (SEND / RECV / DEPT)",
    "       les suivants",
])
corps(
    "Rappel : le temps « saute » d'un événement au suivant. Entre deux événements, "
    "rien ne change, donc inutile de simuler le temps continu.")

h2("3.3  La génération de la trace")
corps(
    "À chaque événement, generateTrace écrit une ligne CSV (time, node, event, "
    "source, destination, msgID ; node 0 = gateway, N = client). Pour les longues "
    "simulations, on désactive la trace (des millions de lignes) ; on en produit "
    "une courte pour illustrer et vérifier l'ordre chronologique. Exemple réel "
    "(le délai de propagation de 1 s se voit entre SEND et RECV) :")
code_bloc([
    "time,node,event,source,destination,msgID",
    "0.328,1,SEND_MSG,1,0,1     ← le client 1 envoie le message 1",
    "1.328,0,RECV_MSG,1,0,1     ← la passerelle le reçoit (Δt = 1,0 s)",
    "1.482,0,MSG_DEPT,1,0,1     ← fin de service du message 1",
])

# 4. Métriques
h1("4.  Les métriques : ce que l'on mesure")
h2("4.1  Nombres et totaux")
corps(
    "On suit le nombre de messages dans le système et dans la file à chaque "
    "instant, et on tient des compteurs totaux : arrivées, services terminés, "
    "messages rejetés.")

h2("4.2  La moyenne temporelle (pondérée par le temps)")
corps(
    "Le nombre de messages dans le système n'est pas constant : il monte et "
    "descend. Pour obtenir une moyenne juste, on pondère par la durée passée dans "
    "chaque état. À chaque événement, on accumule N·Δt (le nombre courant multiplié "
    "par le temps écoulé depuis le dernier événement), puis on divise par T.")
encadre(
    "N̄ = (1/T) · Σ N·Δt . Intuition : passer 9 s avec 1 message et 1 s avec 10 "
    "messages ne donne pas une moyenne de 5,5 mais (9×1 + 1×10)/10 = 1,9. La "
    "pondération par le temps est essentielle.",
    titre="📌  Pourquoi pondérer par le temps ?",
    couleur_fond=(0xFF, 0xF8, 0xE1))

h2("4.3  Les temps moyens d'attente")
corps(
    "On mesure aussi W̄_système (temps moyen passé dans le système, de l'arrivée au "
    "départ) et W̄_file (temps moyen d'attente avant le début du service).")

h2("4.4  La loi de Little (vérification)")
corps(
    "La loi de Little relie ces grandeurs par une formule simple et universelle :")
encadre(
    "L = λ · W   —   le nombre moyen dans le système (L) est égal au taux "
    "d'arrivée (λ) multiplié par le temps moyen de séjour (W). On calcule L et W "
    "de deux façons indépendantes (l'une par accumulation N·Δt, l'autre par les "
    "temps de séjour) : si L ≈ λ·W, c'est une preuve forte que le simulateur est "
    "correct. Dans nos runs, l'égalité est vérifiée partout.",
    titre="✅  L = λ · W : le contrôle du simulateur",
    couleur_fond=(0xE2, 0xEF, 0xDA))

# 5. Modèles
h1("5.  Les modèles simulés")
corps(
    "On simule quatre files qui se distinguent par le nombre de serveurs et la "
    "capacité. La notation M/M/c/K se lit : arrivées Markoviennes / service "
    "Markovien / c serveurs / capacité K (K absent = infinie).")
tableau(
    ["Modèle", "Serveurs", "Capacité K", "Particularité"],
    [["M/M/1", "1", "∞", "file infinie, référence théorique"],
     ["M/M/1/4", "1", "4", "borné : rejets possibles"],
     ["M/M/1/8", "1", "8", "borné : file plus grande"],
     ["M/M/3/8", "3", "8", "3 serveurs, file partagée"]])

h2("5.1  La condition de stabilité")
corps(
    "On définit le taux d'utilisation ρ = λ/µ (par serveur). Une file infinie n'est "
    "stable que si ρ < 1, c'est-à-dire λ < µ : il faut servir plus vite qu'on ne "
    "reçoit. Sinon, la file grossit sans limite.")

h2("5.2  Ce qu'on s'attend à observer")
puce(" : pour ρ < 1 (λ=4, 6) le système est stable ; dès λ ≥ µ = 8, il devient instable (N et W explosent).",
     bold_debut="M/M/1")
puce(" : toujours stables (la borne l'impose), mais des messages sont rejetés ; le rejet augmente avec λ.",
     bold_debut="Modèles bornés")
puce(" : avec 3 serveurs, beaucoup moins de rejets et une attente plus courte qu'un mono-serveur, à charge égale.",
     bold_debut="M/M/3/8")

# 6. Comparaison
h1("6.  Comparaison simulation vs théorie")
corps(
    "Pour M/M/1, la théorie donne des formules fermées : ρ = λ/µ, L = ρ/(1−ρ), "
    "W = 1/(µ−λ). On les utilise comme « vérité de référence » : si la simulation "
    "les retrouve, c'est qu'elle est correcte. (Pour les modèles bornés, on utilise "
    "la formule plus générale M/M/c/K.) Voici nos résultats réels :")
tableau(
    ["Modèle", "λ", "N̄ sim.", "L théo.", "W̄ sim.", "W théo.", "Rejet sim./théo."],
    [["M/M/1", "4", "0,989", "1,000", "0,248", "0,250", "—"],
     ["M/M/1", "6", "2,990", "3,000", "0,498", "0,500", "—"],
     ["M/M/1", "8", "466", "∞", "58,3", "∞", "— (instable)"],
     ["M/M/1/4", "8", "2,000", "2,000", "0,313", "0,313", "19,9 % / 20,0 %"],
     ["M/M/1/8", "12", "6,230", "6,240", "0,788", "0,791", "34,0 % / 34,2 %"],
     ["M/M/3/8", "12", "1,699", "1,706", "0,142", "0,143", "0,37 % / 0,37 %"]])
corps(
    "L'accord est excellent (écart < 1 %). Les cas λ ≥ 8 de M/M/1 « explosent », ce "
    "qui matérialise bien l'instabilité prévue par la théorie (ρ ≥ 1). M/M/3/8 ne "
    "perd presque rien (0,37 %) là où M/M/1/8 perd 34 % à la même charge : "
    "tripler les serveurs change tout.", italic=True)

# ── Récap ──
doc.add_page_break()
h1("Récapitulatif — Séances 2 et 3")
tableau(
    ["Séance", "Classe", "Rôle en une ligne"],
    [["2", "Client", "génère les messages (inter-arrivées exp., moyenne 1/λ)"],
     ["2", "Queue", "file d'attente FIFO, capacité ∞ ou K (rejet si pleine)"],
     ["2", "Server", "traite un message à la fois (service exp., moyenne 1/µ)"],
     ["3", "Gateway", "le système : relie file + serveurs, logique RECV/DEPT"],
     ["3", "Engine", "le moteur : boucle à événements, trace, métriques"]],
    mono_cols=(1,))
encadre(
    "À la fin de la Séance 3, le simulateur est complet : les tests passent, les "
    "simulations tournent, les résultats M/M/1 retrouvent la théorie pour λ=4 et 6, "
    "les modèles bornés montrent des rejets cohérents, et la loi de Little est "
    "vérifiée. Le simulateur est donc validé.",
    titre="✅  Critère de réussite",
    couleur_fond=(0xE2, 0xEF, 0xDA))

# ════════════════════════════════════════════════════════════════════════
chemin = "/home/BaldFox/Documents/TP2/tp_simulateur_java/docs/Seance2-3_Pedagogique.docx"
doc.save(chemin)
print("Document généré :", chemin)
