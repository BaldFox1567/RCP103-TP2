"""
make_doc.py - Génère le document Word pédagogique sur la Séance 1 du TP.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Couleurs ────────────────────────────────────────────────────────────
ROUGE_CNAM  = RGBColor(0xC5, 0x0B, 0x26)
BLEU_TITRE  = RGBColor(0x1F, 0x49, 0x7D)
GRIS_FOND   = RGBColor(0xF2, 0xF2, 0xF2)

doc = Document()

# ── Marges ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Helpers ─────────────────────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, italic=False, color=None, name="Calibri"):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def ajouter_titre_principal(doc, texte):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    set_run_font(run, size=22, bold=True, color=ROUGE_CNAM)
    return p

def ajouter_sous_titre(doc, texte):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    set_run_font(run, size=13, italic=True, color=BLEU_TITRE)
    return p

def h1(doc, texte):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    set_run_font(run, size=15, bold=True, color=ROUGE_CNAM)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # Ligne de séparation via bordure bas
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "C50B26")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(doc, texte):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    set_run_font(run, size=12, bold=True, color=BLEU_TITRE)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def corps(doc, texte, italic=False, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    set_run_font(run, size=11, italic=italic)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    return p

def puce(doc, texte, bold_debut=None, indent=0.8):
    """Paragraphe puce. Si bold_debut fourni, ce préfixe est en gras."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    if bold_debut:
        r1 = p.add_run(bold_debut)
        set_run_font(r1, size=11, bold=True)
        reste = texte
    else:
        reste = texte
    r2 = p.add_run(reste)
    set_run_font(r2, size=11)
    return p

def encadre(doc, texte, titre=None, couleur_fond=None):
    """Paragraphe dans un tableau 1×1 simulant un encadré."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    if couleur_fond:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        r, g, b = couleur_fond
        shd.set(qn("w:fill"),  f"{r:02X}{g:02X}{b:02X}")
        tcPr.append(shd)
    if titre:
        tp = cell.add_paragraph()
        tr = tp.add_run(titre)
        set_run_font(tr, size=11, bold=True, color=ROUGE_CNAM)
        tp.paragraph_format.space_after = Pt(3)
    cp = cell.add_paragraph()
    cr = cp.add_run(texte)
    set_run_font(cr, size=11, italic=False)
    cp.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()   # espace après l'encadré

def code_bloc(doc, lignes):
    """Bloc de code (police Courier, fond gris clair)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), "F2F2F2")
    tcPr.append(shd)
    for ligne in lignes:
        cp = cell.add_paragraph()
        cr = cp.add_run(ligne)
        set_run_font(cr, size=10, name="Courier New")
        cp.paragraph_format.space_after  = Pt(0)
        cp.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()

def saut_ligne(doc):
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ════════════════════════════════════════════════════════════════════════
saut_ligne(doc)
saut_ligne(doc)
ajouter_titre_principal(doc, "Simulateur à événements discrets")
ajouter_sous_titre(doc, "Document pédagogique — Séance 1")
saut_ligne(doc)
ajouter_sous_titre(doc, "RCP103 / NCA — CNAM")
ajouter_sous_titre(doc, "Pedro Braconnot Velloso & Abdelhak Hidouri")
saut_ligne(doc)
encadre(doc,
    "Ce document explique, de façon progressive et illustrée, les notions et "
    "objectifs de la Séance 1 du TP Simulateur. Après lecture, vous saurez "
    "exactement ce que font les classes Message, Event et Scheduler, "
    "pourquoi elles sont conçues ainsi, et comment les tester proprement.",
    titre="📘  À propos de ce document",
    couleur_fond=(0xEB, 0xF3, 0xFF))

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  1. L'IDÉE GÉNÉRALE
# ════════════════════════════════════════════════════════════════════════
h1(doc, "1.  L'idée générale : simuler sans perdre de temps")

h2(doc, "1.1  Pourquoi un simulateur à événements discrets ?")
corps(doc,
    "Imaginons qu'on veuille étudier combien de temps un message attend "
    "avant d'être traité par un serveur. Une approche naïve serait de "
    "simuler chaque milliseconde : \"Est-ce qu'un message arrive à t=0,001 s ? "
    "À t=0,002 s ? À t=0,003 s ?...\" Pour une heure de simulation, ça fait "
    "3 600 000 pas de temps — même si, pendant 99 % de cette durée, il ne "
    "se passe strictement rien.")
encadre(doc,
    "L'idée clé : entre deux événements, rien ne change. Inutile de "
    "simuler ces périodes creuses. On fait donc \"sauter\" le temps "
    "directement d'un événement au suivant. Au lieu de simuler 1 000 000 "
    "d'instants, on ne traite que quelques milliers d'événements — "
    "c'est bien plus rapide.",
    titre="💡  Principe fondamental",
    couleur_fond=(0xFF, 0xF8, 0xE1))

h2(doc, "1.2  Qu'est-ce qu'un événement ?")
corps(doc,
    "Un événement est un fait qui se produit à un instant précis t et qui "
    "modifie l'état du système. Dans notre simulateur, il y a trois types "
    "d'événements :")
puce(doc, " SEND_MSG : un client envoie un message à la passerelle.", bold_debut="SEND_MSG")
puce(doc, " RECV_MSG : la passerelle reçoit le message (après le délai de propagation).", bold_debut="RECV_MSG")
puce(doc, " MSG_DEPT : la passerelle finit de traiter le message (départ).", bold_debut="MSG_DEPT")
saut_ligne(doc)

h2(doc, "1.3  Ce qu'on suit pendant la simulation")
corps(doc, "À chaque événement, le simulateur met à jour plusieurs variables :")

# Tableau des variables
tbl = doc.add_table(rows=4, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
entetes = ["Variable", "Type", "Mise à jour quand ?"]
for i, ent in enumerate(entetes):
    cell = tbl.rows[0].cells[i]
    r = cell.paragraphs[0].add_run(ent)
    set_run_font(r, size=11, bold=True)
    # fond rouge CNAM
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "C50B26")
    tcPr.append(shd)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
lignes_tbl = [
    ("t  (temps courant)", "Horloge", "À chaque événement extrait"),
    ("N_système  /  N_file", "Compteurs d'état", "À chaque arrivée ou départ"),
    ("Nb arrivées, départs, rejets", "Compteurs totaux", "À chaque arrivée ou départ"),
]
for i, (a, b, c) in enumerate(lignes_tbl):
    row = tbl.rows[i+1]
    for j, val in enumerate([a, b, c]):
        r = row.cells[j].paragraphs[0].add_run(val)
        set_run_font(r, size=10.5)
saut_ligne(doc)

# ════════════════════════════════════════════════════════════════════════
#  2. OBJECTIF DE LA SÉANCE 1
# ════════════════════════════════════════════════════════════════════════
h1(doc, "2.  Objectif de la Séance 1")
corps(doc,
    "Le simulateur complet est composé de six classes principales "
    "(Engine, Scheduler, Client, Gateway, Server, Queue) et de deux classes "
    "auxiliaires (Event, Message). On ne construit pas tout d'un coup : "
    "le TP est découpé en trois séances.")
# Tableau des séances
tbl2 = doc.add_table(rows=4, cols=2)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (seance, contenu) in enumerate([
    ("Séance", "Contenu"),
    ("Séance 1  ✔ (maintenant)", "Main/tests, classe Message, classe Event, classe Scheduler"),
    ("Séance 2", "Classes Client, Queue, Server"),
    ("Séance 3", "Classes Engine, Gateway + simulations M/M/1, M/M/1/K, M/M/c/K"),
]):
    r0 = tbl2.rows[i].cells[0].paragraphs[0].add_run(seance)
    r1 = tbl2.rows[i].cells[1].paragraphs[0].add_run(contenu)
    bold = (i == 0 or i == 1)
    set_run_font(r0, size=11, bold=bold)
    set_run_font(r1, size=11, bold=bold)
    if i == 1:
        for cell in tbl2.rows[i].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"),"DAEEF3")
            tcPr.append(shd)
saut_ligne(doc)
encadre(doc,
    "En Séance 1, on pose les trois briques de base sans lesquelles rien "
    "d'autre ne peut fonctionner : un message (ce qu'on transporte), un "
    "événement (ce qui se produit), et le scheduler (qui gère l'ordre des "
    "événements). Tout le reste s'appuiera sur ces trois classes.",
    titre="🎯  Ce qu'on construit aujourd'hui",
    couleur_fond=(0xE2, 0xEF, 0xDA))

# ════════════════════════════════════════════════════════════════════════
#  3. LA CLASSE MESSAGE
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, "3.  La classe Message")

h2(doc, "3.1  À quoi ça sert ?")
corps(doc,
    "Dans notre modèle, un dispositif mobile (le client) envoie des messages "
    "à une passerelle IoT. La classe Message représente un de ces messages : "
    "c'est l'objet de base qui circule dans tout le système. Il est créé par "
    "le client, transporté par les événements, traité par la passerelle.")

h2(doc, "3.2  Ses membres")
# Tableau membres Message
tbl3 = doc.add_table(rows=5, cols=3)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (champ, type_, role) in enumerate([
    ("Champ", "Type", "Signification"),
    ("messageID", "int", "Identifiant unique du message (1, 2, 3…). Permet de retrouver un message dans la trace."),
    ("source", "int", "Nœud émetteur : l'identifiant du client (1, 2…). Utile pour la trace."),
    ("destination", "int", "Nœud destinataire : 0 = la passerelle (convention du sujet)."),
    ("timestamp", "double", "Heure de création du message. Mis à jour par setTimestamp() au moment de l'émission."),
]):
    r0 = tbl3.rows[i].cells[0].paragraphs[0].add_run(champ)
    r1 = tbl3.rows[i].cells[1].paragraphs[0].add_run(type_)
    r2 = tbl3.rows[i].cells[2].paragraphs[0].add_run(role)
    set_run_font(r0, size=10.5, bold=(i==0))
    set_run_font(r1, size=10.5, bold=(i==0), name=("Calibri" if i==0 else "Courier New"))
    set_run_font(r2, size=10.5, bold=(i==0))
    if i == 0:
        for cell in tbl3.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"),"C50B26")
            tcPr.append(shd)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
saut_ligne(doc)

h2(doc, "3.3  Ses méthodes")
puce(doc, "Message(id, src, dst)  — constructeur. Le timestamp est initialisé à 0.", bold_debut="Constructeur : ")
puce(doc, "getMessageID(), getSource(), getDestination(), getTimestamp()  — lecture des champs.", bold_debut="Getters : ")
puce(doc, "setTimestamp(t)  — met à jour l'heure de création (appelé au moment de l'émission).", bold_debut="Setter : ")
puce(doc, "printMessage()  — affiche tous les membres. Utile pour déboguer.", bold_debut="Affichage : ")
saut_ligne(doc)

corps(doc, "Exemple de sortie de printMessage() :")
code_bloc(doc, [
    "Message #1 : source=1 -> destination=0, timestamp=1.202"
])
encadre(doc,
    "Le timestamp du message mémorise l'instant où le message a été créé. "
    "Plus tard, quand le message quittera le système (MSG_DEPT à l'instant t_depart), "
    "on calculera le temps de séjour = t_depart - timestamp. "
    "C'est la métrique W (temps moyen dans le système) du cours.",
    titre="📌  Pourquoi le timestamp est-il important ?",
    couleur_fond=(0xFF, 0xF8, 0xE1))

# ════════════════════════════════════════════════════════════════════════
#  4. LA CLASSE EVENT
# ════════════════════════════════════════════════════════════════════════
h1(doc, "4.  La classe Event (événement)")

h2(doc, "4.1  Qu'est-ce qu'un événement, concrètement ?")
corps(doc,
    "Un Event est une \"fiche\" qui dit : \"À l'instant t, telle chose va se "
    "passer pour tel message.\" Le Scheduler stocke ces fiches dans l'ordre "
    "chronologique et les distribue une par une à la boucle de simulation.")

h2(doc, "4.2  Ses membres")
tbl4 = doc.add_table(rows=5, cols=3)
tbl4.style = "Table Grid"
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (champ, type_, role) in enumerate([
    ("Champ", "Type", "Signification"),
    ("eventID", "int", "Identifiant unique de l'événement. Permet de le distinguer dans la liste du Scheduler."),
    ("message", "Message", "Le message concerné par cet événement (référence vers l'objet Message)."),
    ("eventType", "EventType", "Le type : SEND_MSG, RECV_MSG ou MSG_DEPT (voir ci-dessous)."),
    ("eventTime", "double", "L'instant auquel l'événement est programmé (en secondes depuis t=0)."),
]):
    r0 = tbl4.rows[i].cells[0].paragraphs[0].add_run(champ)
    r1 = tbl4.rows[i].cells[1].paragraphs[0].add_run(type_)
    r2 = tbl4.rows[i].cells[2].paragraphs[0].add_run(role)
    set_run_font(r0, size=10.5, bold=(i==0))
    set_run_font(r1, size=10.5, bold=(i==0), name=("Calibri" if i==0 else "Courier New"))
    set_run_font(r2, size=10.5, bold=(i==0))
    if i == 0:
        for cell in tbl4.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"),"C50B26")
            tcPr.append(shd)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
saut_ligne(doc)

h2(doc, "4.3  Les trois types d'événements (enum EventType)")
corps(doc,
    "On représente les types par une énumération (enum), ce qui évite les "
    "erreurs de frappe et rend le code lisible :")
code_bloc(doc, [
    "public enum EventType {",
    "    SEND_MSG,   // le client émet un message",
    "    RECV_MSG,   // la passerelle reçoit le message",
    "    MSG_DEPT    // fin de service : le message quitte le système",
    "}",
])
corps(doc, "Ce que chaque type provoque dans la simulation :")
puce(doc,
    " : le client émet. On programme un RECV_MSG pour dans 1 seconde "
    "(délai de propagation), et on programme le prochain SEND_MSG du même client.",
    bold_debut="SEND_MSG")
puce(doc,
    " : la passerelle reçoit. Si un serveur est libre → service immédiat + "
    "programmer MSG_DEPT. Sinon → mise en file d'attente. Si système plein (cas /K) → rejet.",
    bold_debut="RECV_MSG")
puce(doc,
    " : fin de service. Libérer le serveur. S'il y a des messages en file → "
    "servir le suivant.",
    bold_debut="MSG_DEPT")
saut_ligne(doc)

h2(doc, "4.4  Ses méthodes")
puce(doc, "Event(id, type, time, msg)  — constructeur.", bold_debut="Constructeur : ")
puce(doc, "setEventTime(t) / getEventTime()  — accès à l'horodatage.", bold_debut="Temps : ")
puce(doc, "setEventType(t) / getEventType()  — accès au type.", bold_debut="Type : ")
puce(doc, "printEvent()  — affiche tous les membres. Ex. : Event #4 : type=RECV_MSG  t=1.916  msg=1", bold_debut="Affichage : ")
saut_ligne(doc)

encadre(doc,
    "En Java, Event doit implémenter Comparable<Event>. La méthode compareTo() "
    "compare les eventTime, et en cas d'égalité, les eventID. Cela permet à la "
    "PriorityQueue du Scheduler de trier automatiquement les événements — "
    "sans que le Scheduler ait besoin de connaître les détails d'un Event.",
    titre="🔧  Détail d'implémentation : Comparable<Event>",
    couleur_fond=(0xF2, 0xF2, 0xF2))

# ════════════════════════════════════════════════════════════════════════
#  5. LA CLASSE SCHEDULER
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, "5.  La classe Scheduler (ordonnanceur)")

h2(doc, "5.1  Son rôle")
corps(doc,
    "Le Scheduler est le cœur du simulateur à événements discrets. "
    "Il maintient la liste de tous les événements futurs, TOUJOURS dans "
    "l'ordre chronologique. La boucle de simulation ne fait qu'extraire "
    "le premier événement (le plus proche dans le temps), le traiter, "
    "puis en ajouter de nouveaux.")

h2(doc, "5.2  Ses méthodes")
tbl5 = doc.add_table(rows=5, cols=2)
tbl5.style = "Table Grid"
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (meth, role) in enumerate([
    ("Méthode", "Ce qu'elle fait"),
    ("addEvent(e)", "Insère l'événement e à la bonne position (ordre chronologique maintenu)."),
    ("getEvent()", "Extrait et retourne le prochain événement (le plus tôt). Met à jour t."),
    ("getCurrentTime()", "Retourne le temps courant t = temps du dernier événement extrait."),
    ("hasEvents()", "Retourne true s'il reste des événements à traiter (boucle non terminée)."),
]):
    r0 = tbl5.rows[i].cells[0].paragraphs[0].add_run(meth)
    r1 = tbl5.rows[i].cells[1].paragraphs[0].add_run(role)
    bold = (i == 0)
    set_run_font(r0, size=10.5, bold=bold, name=("Calibri" if bold else "Courier New"))
    set_run_font(r1, size=10.5, bold=bold)
    if bold:
        for cell in tbl5.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"),"C50B26")
            tcPr.append(shd)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
saut_ligne(doc)

h2(doc, "5.3  Pourquoi on n'utilise pas sort() ?")
corps(doc,
    "Imaginons qu'on stocke les événements dans une liste ordinaire et qu'on "
    "appelle sort() après chaque addEvent(). Problème : sort() re-trie toute "
    "la liste à chaque fois — coût O(n log n) par insertion. Pour une "
    "simulation qui génère des millions d'événements, c'est catastrophique.")
corps(doc,
    "La solution : utiliser une structure qui maintient l'ordre "
    "automatiquement à chaque insertion.")

# Tableau comparaison
tbl6 = doc.add_table(rows=4, cols=3)
tbl6.style = "Table Grid"
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (struct, insert, extract) in enumerate([
    ("Structure", "Insertion", "Extraction du min"),
    ("Liste + sort()", "O(n log n)  ❌ lent", "O(1)"),
    ("Liste + insertion ordonnée", "O(n)  ⚠ moyen", "O(1)"),
    ("PriorityQueue (tas binaire)", "O(log n)  ✅ rapide", "O(log n)  ✅ rapide"),
]):
    for j, val in enumerate([struct, insert, extract]):
        r = tbl6.rows[i].cells[j].paragraphs[0].add_run(val)
        set_run_font(r, size=10.5, bold=(i==0))
    if i == 0:
        for cell in tbl6.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"),"C50B26")
            tcPr.append(shd)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    if i == 3:
        for cell in tbl6.rows[3].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"),"E2EFDA")
            tcPr.append(shd)
saut_ligne(doc)

encadre(doc,
    "On utilise java.util.PriorityQueue (tas binaire min-heap). "
    "offer(e) insère en O(log n) : le tas se réorganise seul. "
    "poll() extrait toujours le minimum (l'événement le plus tôt) en O(log n). "
    "Il n'y a jamais de sort() appelé. L'ordre chronologique est garanti "
    "par la méthode compareTo() de la classe Event.",
    titre="✅  Notre choix : PriorityQueue",
    couleur_fond=(0xE2, 0xEF, 0xDA))

h2(doc, "5.4  Illustration : le Scheduler en action")
corps(doc,
    "On insère 4 événements dans le désordre. Le Scheduler les restitue "
    "toujours dans l'ordre chronologique :")
code_bloc(doc, [
    "// Insertion dans le désordre :",
    "scheduler.addEvent(new Event(1, MSG_DEPT,  4.572, m1));",
    "scheduler.addEvent(new Event(2, SEND_MSG,  1.202, m1));",
    "scheduler.addEvent(new Event(3, SEND_MSG,  2.320, m2));",
    "scheduler.addEvent(new Event(4, RECV_MSG,  1.916, m1));",
    "",
    "// Extraction en ordre chronologique :",
    "→ t=1.202  SEND_MSG  msg=1",
    "→ t=1.916  RECV_MSG  msg=1",
    "→ t=2.320  SEND_MSG  msg=2",
    "→ t=4.572  MSG_DEPT  msg=1",
])
corps(doc,
    "Ces quatre instants (1.202, 1.916, 2.320, 4.572) correspondent exactement "
    "à la trace d'exemple présentée dans les diapositives du cours (diapo 13).",
    italic=True)

# ════════════════════════════════════════════════════════════════════════
#  6. LA TRACE
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, "6.  La trace : enregistrement des événements")

h2(doc, "6.1  À quoi sert la trace ?")
corps(doc,
    "La trace est un journal chronologique de tous les événements qui se "
    "sont produits pendant la simulation. Elle sert à :")
puce(doc, "Déboguer : vérifier que l'ordre chronologique est respecté, que les bons messages arrivent aux bons nœuds.")
puce(doc, "Calculer les métriques a posteriori : temps d'arrivée, temps de départ, temps de séjour.")
puce(doc, "Visualiser le comportement du système : voir comment les files se remplissent.")
saut_ligne(doc)

h2(doc, "6.2  Format de la trace (colonnes)")
corps(doc, "Après chaque événement, on écrit une ligne CSV avec les colonnes suivantes :")
tbl7 = doc.add_table(rows=7, cols=2)
tbl7.style = "Table Grid"
tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (col, signif) in enumerate([
    ("Colonne", "Signification"),
    ("time", "Temps de l'événement (en secondes, 3 décimales)."),
    ("node", "Nœud concerné : 0 = passerelle (gateway), N = client N."),
    ("event", "Type de l'événement : SEND_MSG, RECV_MSG ou MSG_DEPT."),
    ("source", "Identifiant du nœud émetteur du message."),
    ("destination", "Identifiant du nœud destinataire (0 = passerelle)."),
    ("msgID", "Identifiant unique du message concerné."),
]):
    r0 = tbl7.rows[i].cells[0].paragraphs[0].add_run(col)
    r1 = tbl7.rows[i].cells[1].paragraphs[0].add_run(signif)
    set_run_font(r0, size=10.5, bold=(i==0), name=("Calibri" if i==0 else "Courier New"))
    set_run_font(r1, size=10.5, bold=(i==0))
    if i == 0:
        for cell in tbl7.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"),"C50B26")
            tcPr.append(shd)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
saut_ligne(doc)

h2(doc, "6.3  Exemple commenté (trace réelle du simulateur)")
corps(doc, "Voici les premières lignes d'une trace réelle (M/M/1, λ=4) :")
code_bloc(doc, [
    "time,node,event,source,destination,msgID",
    "1.202,1,SEND_MSG,1,0,1   ← t=1.202 : le client 1 envoie le message 1",
    "1.916,0,RECV_MSG,1,0,1   ← t=1.916 : la passerelle reçoit msg 1  (Δt = 0.714 s de propagation)",
    "2.320,3,SEND_MSG,3,0,2   ← t=2.320 : le client 3 envoie le message 2",
    "2.391,0,RECV_MSG,3,0,2   ← t=2.391 : la passerelle reçoit msg 2",
    "4.572,0,MSG_DEPT,1,0,1   ← t=4.572 : service du msg 1 terminé (séjour = 4.572 - 1.916 = 2.656 s)",
    "5.916,0,MSG_DEPT,3,0,2   ← t=5.916 : service du msg 2 terminé",
])
corps(doc,
    "On vérifie que les temps sont bien croissants (ordre chronologique). "
    "C'est exactement ce que testScheduler() doit vérifier automatiquement.",
    italic=True)

h2(doc, "6.4  La méthode generateTrace()")
corps(doc,
    "Dans le Main (et plus tard dans l'Engine), on appelle generateTrace(e, node) "
    "après chaque traitement d'événement. Elle affiche ou écrit la ligne CSV :")
code_bloc(doc, [
    "public static void generateTrace(Event e, int node) {",
    "    Message m = e.getMessage();",
    "    System.out.printf(Locale.ROOT,",
    "        \"%.3f,%d,%s,%d,%d,%d%n\",",
    "        e.getEventTime(), node, e.getEventType(),",
    "        m.getSource(), m.getDestination(), m.getMessageID());",
    "}",
])

# ════════════════════════════════════════════════════════════════════════
#  7. MÉTHODOLOGIE DE TEST
# ════════════════════════════════════════════════════════════════════════
h1(doc, "7.  Méthodologie de test")

h2(doc, "7.1  Principe du développement incrémental")
encadre(doc,
    "On implémente et on teste UNE classe à la fois. On ne passe à la suivante "
    "que quand la précédente est validée. Ici : Message → Event → Scheduler. "
    "Si une erreur apparaît, on sait dans quelle classe elle se trouve.",
    titre="📋  Règle d'or",
    couleur_fond=(0xEB, 0xF3, 0xFF))

h2(doc, "7.2  Une méthode Print par classe")
corps(doc,
    "Chaque classe a une méthode d'affichage (printMessage, printEvent) "
    "qui affiche tous ses membres. C'est le premier test : si l'affichage "
    "est correct, les membres sont bien initialisés.")

h2(doc, "7.3  Une fonction de test par classe dans le Main")
corps(doc, "On écrit une fonction statique dans Main pour chaque classe :")
puce(doc,
    " : crée un Message(1, 1, 0), appelle setTimestamp(1.202), "
    "printMessage(), puis vérifie avec des assert que getMessageID()==1, "
    "getTimestamp()==1.202, etc.",
    bold_debut="testMessage()")
puce(doc,
    " : crée un Event, modifie son type et son temps, printEvent(), "
    "appelle generateTrace() pour afficher la ligne CSV.",
    bold_debut="testEvent()")
puce(doc,
    " : insère plusieurs événements dans le désordre, appelle getEvent() "
    "en boucle (hasEvents()), vérifie que les temps ressortent bien en "
    "ordre croissant.",
    bold_debut="testScheduler()")
saut_ligne(doc)

corps(doc, "Exemple : ce que testScheduler() vérifie :")
code_bloc(doc, [
    "// Insérer 4 événements dans le désordre...",
    "scheduler.addEvent(new Event(1, MSG_DEPT,  4.572, m1));",
    "scheduler.addEvent(new Event(2, SEND_MSG,  1.202, m1));",
    "// ...",
    "",
    "// Vérifier qu'ils ressortent dans l'ordre :",
    "double tPrecedent = -1.0;",
    "while (scheduler.hasEvents()) {",
    "    Event e = scheduler.getEvent();",
    "    assert e.getEventTime() >= tPrecedent;  // ← doit toujours être vrai",
    "    tPrecedent = e.getEventTime();",
    "}",
])

h2(doc, "7.4  Activer les assertions Java")
corps(doc,
    "En Java, les assert ne s'exécutent que si on lance la JVM avec le flag -ea "
    "(Enable Assertions). Sans ce flag, les assert sont ignorés silencieusement.")
code_bloc(doc, [
    "javac -encoding UTF-8 *.java   # compilation",
    "java -ea Main                   # exécution avec assertions activées",
])
encadre(doc,
    "Si un assert échoue, Java lève une AssertionError avec le message qu'on a "
    "fourni. Exemple : AssertionError: ERREUR : ordre non respecté ! t=1.202 < 1.916. "
    "C'est beaucoup plus informatif qu'un simple crash.",
    titre="💡  assert en Java",
    couleur_fond=(0xFF, 0xF8, 0xE1))

# ════════════════════════════════════════════════════════════════════════
#  RÉCAP FINAL
# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, "Récapitulatif — Ce qu'on doit avoir à la fin de la Séance 1")

tbl8 = doc.add_table(rows=6, cols=3)
tbl8.style = "Table Grid"
tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (fichier, contenu, valide) in enumerate([
    ("Fichier", "Contenu", "Validé par"),
    ("EventType.java", "enum SEND_MSG, RECV_MSG, MSG_DEPT", "—"),
    ("Message.java", "messageID, source, destination, timestamp + méthodes", "testMessage()"),
    ("Event.java", "eventID, message, eventType, eventTime + Comparable", "testEvent()"),
    ("Scheduler.java", "PriorityQueue + addEvent, getEvent, getCurrentTime, hasEvents", "testScheduler()"),
    ("Main.java", "testMessage(), testEvent(), testScheduler(), generateTrace()", "java -ea Main"),
]):
    r0 = tbl8.rows[i].cells[0].paragraphs[0].add_run(fichier)
    r1 = tbl8.rows[i].cells[1].paragraphs[0].add_run(contenu)
    r2 = tbl8.rows[i].cells[2].paragraphs[0].add_run(valide)
    set_run_font(r0, size=10.5, bold=(i==0), name=("Calibri" if i==0 else "Courier New"))
    set_run_font(r1, size=10.5, bold=(i==0))
    set_run_font(r2, size=10.5, bold=(i==0), name=("Calibri" if i==0 else "Courier New"))
    if i == 0:
        for cell in tbl8.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
            shd.set(qn("w:fill"),"C50B26")
            tcPr.append(shd)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
saut_ligne(doc)

encadre(doc,
    "À la fin de la séance, java -ea Main doit afficher les trois blocs de test "
    "sans aucun AssertionError, et la trace CSV doit montrer des temps "
    "strictement croissants. Si c'est le cas, les bases du simulateur "
    "sont solides et on peut passer à la Séance 2.",
    titre="✅  Critère de réussite de la Séance 1",
    couleur_fond=(0xE2, 0xEF, 0xDA))

# ════════════════════════════════════════════════════════════════════════
#  SAUVEGARDE
# ════════════════════════════════════════════════════════════════════════
chemin = "/home/BaldFox/Documents/TP2/tp_simulateur_java/docs/Seance1_Pedagogique.docx"
doc.save(chemin)
print(f"Document généré : {chemin}")
